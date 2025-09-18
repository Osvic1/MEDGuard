from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from datetime import datetime, date, timedelta
from flask import Blueprint, request, send_file, jsonify, url_for, render_template, Response
from sqlite3 import IntegrityError
from backend.models import insert_drug
from backend.database import get_db
import qrcode
import io
import traceback
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

admin_bp = Blueprint("admin_api", __name__)

# =========================
# Admin Dashboard
# =========================


@admin_bp.route('/admin')
def admin_dashboard():
    try:
        conn = get_db()
        # Fetch the last 5 reports as a preview
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            ORDER BY reported_on DESC
            LIMIT 5
        """).fetchall()
        reports = [dict(r) for r in rows]
        return render_template('admin.html', reports=reports)
    except Exception as e:
        print("Error in /admin:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Register a new drug batch
# =========================


@admin_bp.post("/register")
def admin_register():
    try:
        data = request.get_json(silent=True) or request.form.to_dict()

        required = ["name", "batch_number",
                    "mfg_date", "expiry_date", "manufacturer"]
        missing = [f for f in required if not data.get(f)]
        if missing:
            return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

        try:
            insert_drug(
                data["name"],
                data["batch_number"],
                data["mfg_date"],
                data["expiry_date"],
                data["manufacturer"]
            )
        except IntegrityError:
            return jsonify({"error": "Batch number already exists"}), 409

        # Encode the correct relative URL for the verify endpoint
        verify_url = f"/verify/{data['batch_number']}"

        # Generate QR with increased box size and border for better scanning
        qr_img = qrcode.make(verify_url, box_size=15, border=4)
        buf = io.BytesIO()
        qr_img.save(buf, format="PNG")
        buf.seek(0)

        if request.is_json:
            return send_file(buf, mimetype="image/png", download_name=f"{data['batch_number']}.png")

        import base64
        qr_base64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        # Fetch latest reports
        conn = get_db()
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            ORDER BY reported_on DESC
            LIMIT 5
        """).fetchall()
        reports = [dict(r) for r in rows]

        return render_template("admin.html", qr_image=qr_base64, reports=reports, scroll='qr')

    except Exception as e:
        print("Error in /register:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Registered Drugs Page
# =========================


@admin_bp.get("/drugs")
def admin_drugs():
    try:
        conn = get_db()
        search = request.args.get("search", "").strip()
        status = request.args.get(
            "status", "").strip()   # valid, expired, soon
        start = request.args.get("start", "").strip()     # 👈 NEW
        end = request.args.get("end", "").strip()         # 👈 NEW
        page = int(request.args.get("page", 1))
        per_page = 20
        offset = (page - 1) * per_page

        base_query = "FROM drugs WHERE 1=1"
        params = []

        # Search filter
        if search:
            base_query += " AND (name LIKE ? OR batch_number LIKE ?)"
            params.extend([f"%{search}%", f"%{search}%"])

        # Expiry filter
        if status == "valid":
            base_query += " AND date(expiry_date) >= date('now','localtime')"
        elif status == "expired":
            base_query += " AND date(expiry_date) < date('now','localtime')"
        elif status == "soon":
            base_query += " AND date(expiry_date) BETWEEN date('now','localtime') AND date('now','+30 day','localtime')"

        # Date range filter
        if start and end:
            base_query += " AND date(created_at) BETWEEN date(?) AND date(?)"
            params.extend([start, end])

        # Count + fetch
        total = conn.execute(
            f"SELECT COUNT(*) {base_query}", params).fetchone()[0]
        rows = conn.execute(f"""
            SELECT name, batch_number, manufacturer, mfg_date, expiry_date, created_at
            {base_query}
            ORDER BY created_at DESC
            LIMIT ? OFFSET ?
        """, params + [per_page, offset]).fetchall()

        total_pages = (total + per_page - 1) // per_page

        return render_template(
            "admin_drugs.html",
            drugs=rows,
            search=search,
            status=status,
            start=start,   # 👈 pass to template
            end=end,       # 👈 pass to template
            current_date=date.today().isoformat(),
            soon_date=(date.today() + timedelta(days=30)).isoformat(),
            page=page,
            total_pages=total_pages
        )
    except Exception as e:
        print("Error in /drugs:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Export Registered Drugs (Word)
# =========================


@admin_bp.get("/drugs/export/word")
def export_drugs_word():
    try:
        conn = get_db()
        search = request.args.get("search", "").strip()
        status = request.args.get("status", "").strip()
        start = request.args.get("start", "").strip()
        end = request.args.get("end", "").strip()

        base_query = "FROM drugs WHERE 1=1"
        params = []

        if search:
            base_query += " AND (name LIKE ? OR batch_number LIKE ?)"
            params.extend([f"%{search}%", f"%{search}%"])

        if status == "valid":
            base_query += " AND date(expiry_date) >= date('now','localtime')"
        elif status == "expired":
            base_query += " AND date(expiry_date) < date('now','localtime')"
        elif status == "soon":
            base_query += " AND date(expiry_date) BETWEEN date('now','localtime') AND date('now','+30 day','localtime')"

        if start and end:
            base_query += " AND date(created_at) BETWEEN date(?) AND date(?)"
            params.extend([start, end])

        rows = conn.execute(f"""
            SELECT name, batch_number, manufacturer, mfg_date, expiry_date, created_at
            {base_query}
            ORDER BY created_at DESC
        """, params).fetchall()

        today = date.today()
        soon = today + timedelta(days=30)

        def safe(val):
            return str(val) if val is not None else "N/A"

        doc = Document()
        doc.add_heading("Registered Drugs Report", 0)

        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Name"
        hdr_cells[1].text = "Batch Number"
        hdr_cells[2].text = "Manufacturer"
        hdr_cells[3].text = "Mfg Date"
        hdr_cells[4].text = "Expiry Date"
        hdr_cells[5].text = "Status"
        hdr_cells[6].text = "Registered On"

        for row in rows:
            try:
                exp_date = datetime.strptime(
                    safe(row["expiry_date"]), "%Y-%m-%d").date()
            except Exception:
                exp_date = today

            if exp_date < today:
                status_label = "Expired"
            elif exp_date <= soon:
                status_label = "Expiring Soon"
            else:
                status_label = "Valid"

            cells = table.add_row().cells
            cells[0].text = safe(row["name"])
            cells[1].text = safe(row["batch_number"])
            cells[2].text = safe(row["manufacturer"])
            cells[3].text = safe(row["mfg_date"])
            cells[4].text = safe(row["expiry_date"])
            cells[5].text = status_label
            cells[6].text = safe(row["created_at"])

        doc.add_paragraph()
        generated_on = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on {generated_on} by Admin System")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name="registered_drugs.docx",
            as_attachment=True
        )
    except Exception as e:
        print("Error exporting Word:", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# =========================
# Export Registered Drugs (PDF)
# =========================


@admin_bp.get("/drugs/export/pdf")
def export_drugs_pdf():
    try:
        conn = get_db()
        search = request.args.get("search", "").strip()
        status = request.args.get("status", "").strip()
        start = request.args.get("start", "").strip()
        end = request.args.get("end", "").strip()

        base_query = "FROM drugs WHERE 1=1"
        params = []

        if search:
            base_query += " AND (name LIKE ? OR batch_number LIKE ?)"
            params.extend([f"%{search}%", f"%{search}%"])

        if status == "valid":
            base_query += " AND date(expiry_date) >= date('now','localtime')"
        elif status == "expired":
            base_query += " AND date(expiry_date) < date('now','localtime')"
        elif status == "soon":
            base_query += " AND date(expiry_date) BETWEEN date('now','localtime') AND date('now','+30 day','localtime')"

        if start and end:
            base_query += " AND date(created_at) BETWEEN date(?) AND date(?)"
            params.extend([start, end])

        rows = conn.execute(f"""
            SELECT name, batch_number, manufacturer, mfg_date, expiry_date, created_at
            {base_query}
            ORDER BY created_at DESC
        """, params).fetchall()

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        # Title
        elements.append(Paragraph("Registered Drugs Report", styles["Title"]))
        elements.append(Spacer(1, 12))

        # Table data with Status
        data = [["Name", "Batch Number", "Manufacturer",
                 "Mfg Date", "Expiry Date", "Status", "Registered On"]]
        today = date.today()
        soon = today + timedelta(days=30)

        for row in rows:
            if row["expiry_date"] < today.isoformat():
                status_label = "Expired"
            elif row["expiry_date"] <= soon.isoformat():
                status_label = "Expiring Soon"
            else:
                status_label = "Valid"

            data.append([
                row["name"], row["batch_number"], row["manufacturer"],
                row["mfg_date"], row["expiry_date"], status_label, row["created_at"]
            ])

        # Styled table
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f4f4f4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 11),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BOX", (0, 0), (-1, -1), 1, colors.black),
        ]))

        elements.append(table)
        elements.append(Spacer(1, 20))

        # Footer
        generated_on = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        footer_text = f"Generated on {generated_on} by Admin System"
        elements.append(Paragraph(footer_text, styles["Normal"]))

        doc.build(elements)
        buf.seek(0)

        return send_file(buf, mimetype="application/pdf",
                         download_name="registered_drugs.pdf", as_attachment=True)
    except Exception as e:
        print("Error exporting PDF:", e)
        traceback.print_exc()
        return jsonify({"error": "Failed to export PDF"}), 500

# =========================
# Reports (All)
# =========================


@admin_bp.get("/reports")
def admin_reports():
    try:
        conn = get_db()
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            ORDER BY reported_on DESC, id DESC
        """).fetchall()
        conn.execute("UPDATE reports SET status = 1 WHERE status = 0")
        conn.commit()
        return render_template("admin.html", reports=rows, scroll='reports')
    except Exception as e:
        print("Error in /reports:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Reports (Today Only)
# =========================


@admin_bp.get("/reports/today")
def admin_reports_today():
    try:
        conn = get_db()
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            WHERE date(reported_on) = date('now','localtime')
            ORDER BY reported_on DESC, id DESC
        """).fetchall()
        return render_template("admin.html", reports=rows, scroll='reports')
    except Exception as e:
        print("Error in /reports/today:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Reports by Date Range
# =========================


@admin_bp.get("/reports/range")
def admin_reports_range():
    try:
        start = request.args.get("start")
        end = request.args.get("end")

        if not start or not end:
            return jsonify({"error": "Please provide start and end dates (YYYY-MM-DD)"}), 400

        conn = get_db()
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            WHERE date(reported_on) BETWEEN date(?) AND date(?)
            ORDER BY reported_on DESC, id DESC
        """, (start, end)).fetchall()

        return render_template("admin.html", reports=rows, scroll='reports')
    except Exception as e:
        print("Error in /reports/range:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Reports Count (Unread Today Only)
# =========================


@admin_bp.get("/reports/count")
def reports_count():
    try:
        conn = get_db()
        row = conn.execute("""
            SELECT COUNT(*) AS total
            FROM reports
            WHERE status = 0
              AND date(reported_on) = date('now','localtime')
        """).fetchone()
        return jsonify({"count": row["total"]})
    except Exception as e:
        print("Error in /reports/count:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# =========================
# Reports Preview (Today Only)
# =========================


@admin_bp.get("/reports/preview")
def reports_preview():
    try:
        conn = get_db()
        rows = conn.execute("""
            SELECT id, drug_name, batch_number, location, note,
                   strftime('%Y-%m-%d %H:%M:%S', reported_on) AS reported_on, status
            FROM reports
            WHERE date(reported_on) = date('now','localtime')
            ORDER BY reported_on DESC
            LIMIT 5
        """).fetchall()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        print("Error in /reports/preview:", e)
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500
